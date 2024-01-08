import streamlit as st
from openai import OpenAI
from docx import Document
import io
import os
import datetime
import geopandas as gpd

# Function to load countries data
def load_countries():
    data = gpd.read_file(os.path.join('data', 'merged_file.gpkg'))
    data = data[data['field_3'].notna()]
    return sorted(data['field_3'].unique().tolist())

# Initialize session state variables if they don't exist
if 'selected_countries' not in st.session_state:
    st.session_state['selected_countries'] = []
if 'participants' not in st.session_state:
    st.session_state['participants'] = []
if 'messages' not in st.session_state:
    st.session_state['messages'] = []
    
openai_api_key = os.getenv('openai_api_key')

# Page 1: Initial Setup
def page1():
    with st.sidebar:
        # GPT Selection
        gpt_version = st.radio("Choose GPT Version", ('3.5', '4'))
        
        # Travel Dates
        start_date = st.date_input("Start Date",value=datetime.date(2023, 12, 7))
        end_date = st.date_input("End Date",value=datetime.date(2024, 1, 7))

        # Country Selection
        countries = load_countries()
        selected_countries = st.multiselect('Choose countries', countries, default=['Thailand','Laos'])
        st.session_state.selected_countries = selected_countries

        # Number of People
        num_people = st.number_input("Number of People", min_value=1, max_value=100)

    # Save settings to session state
    st.session_state['gpt_version'] = gpt_version
    st.session_state['start_date'] = start_date
    st.session_state['end_date'] = end_date
    st.session_state['num_people'] = num_people

# Page 2: Participant Details
def page2():
    st.session_state['participants'] = []
    for i in range(st.session_state['num_people']):
        st.subheader(f"Participant {i+1}")
        with st.form(f"participant_{i}"):
            name = st.text_input(f"Name of Participant {i+1}", value="Darlain")
            age = st.number_input(f"Age of Participant {i+1}", min_value=0, max_value=120,value=30)
            gender = st.selectbox(f"Gender of Participant {i+1}", ['Male', 'Female', 'Other'])
            preference = st.selectbox(f"Vacation Preference of Participant {i+1}", ['Adventure', 'Relax', 'Culture'])
            additional_preferences = st.text_area("Additional Preferences",value="I like to discover new cultures")
            submitted = st.form_submit_button("Save Participant")

            if submitted:
                participant = {
                    'name': name,
                    'age': age,
                    'gender': gender,
                    'preference': preference,
                    'additional_preferences': additional_preferences
                }
                st.session_state['participants'].append(participant)

# Page 3: Dynamic Chatbot for Travel Planning
def page3():
    st.title("💬 Dynamic Travel Planning Chatbot")

    # Initialize chatbot with specific context-based questions
    if "chat_initialized" not in st.session_state:
        st.session_state.chat_initialized = True
        travel_context = generate_travel_context()
        initialize_chat_with_context(travel_context)
        
    # Display chatbot messages
    display_chatbot_messages()

    # User input for chatbot
    handle_user_input()


# Helper function to get the correct chatbot model
def get_chatbot_model():
    return "gpt-3.5-turbo" if st.session_state['gpt_version'] == '3.5' else "gpt-4"
    
# Function to generate the travel context for the chatbot
def generate_travel_context():
    context = "As a travel agent, I need to refine our travel plan. Here's the information I have:\n"
    context += f"Travel Dates: {st.session_state['start_date']} to {st.session_state['end_date']}\n"
    context += f"Destinations: {', '.join(st.session_state['selected_countries'])}\n"
    context += "Only participant(s):\n"
    for participant in st.session_state['participants']:
        context += f"- {participant['name']}, {participant['age']} years old, {participant['gender']}, "
        context += f"prefers {participant['preference']}. Additional notes: {participant['additional_preferences']}\n"
    return context
    
# Function to initialize chatbot with travel context
def initialize_chat_with_context(travel_context):
    context_prompt = f"You are a travel agent and based on the following travel plan details:\n{travel_context}\nGenerate one question to refine the trip planning until you think you have a good plan to suggest. You are talking with the traveller(s). Please start from the assumption that they don't know nothing about the countries they are visiting so instead of asking for suggestions about names of places they would like to go, try to propose options for them to choose so that you learn their tastes. Take into account the total time they said they want they trip to last (you have the info above), and make sure you help them plan the best trip ever."
    st.session_state.messages = [{"role": "system", "content": context_prompt}]
    generate_next_question()

# Function to display chatbot messages
def display_chatbot_messages():
    for msg in st.session_state.messages:
        st.chat_message(msg["role"]).write(msg["content"])

# Function to handle user input and generate next question
def handle_user_input():
    if prompt := st.chat_input():
        # Append user input to messages and write it immediately
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.chat_message("user").write(prompt)

        # Generate next question (chatbot response)
        generate_next_question()

        # Write the chatbot response immediately
        last_message = st.session_state.messages[-1]
        st.chat_message(last_message["role"]).write(last_message["content"])


# Function to generate next question after user input
def generate_next_question():
    client = OpenAI(api_key=openai_api_key)
    next_question_response = client.chat.completions.create(
        model=get_chatbot_model(),
        messages=st.session_state.messages
    )
    next_question = next_question_response.choices[0].message.content
    st.session_state.messages.append({"role": "assistant", "content": next_question})
    
def page4():
    st.title("🌍 Final Trip Overview")

    # Button to generate itinerary
    if st.sidebar.button("Generate Itinerary"):
        itinerary = generate_itinerary_from_conversation(st.session_state.messages)
        st.subheader("Your Customized Travel Plan:")
        st.write(itinerary)

def generate_itinerary_from_conversation(messages):
    # Format the conversation into a prompt
    prompt = "\n".join([f"{msg['role']}: {msg['content']}" for msg in messages])
    
    # Add additional instructions for the AI if needed
    prompt += "\n\nBased on the above conversation, generate the best travel itinerary. Take into account that this conversation is not meant to include everything that should be in the plan. It helps you understanding the tastes of the travellers. Use it as a reference but then propose them the plan you think would be best, based on available time and optimizing logistics etc. It has to be a day by day plan, that takes into account feasibility, travel, movements, tastes of the participants, and includes the must see places in the country or countries, the time of the year. You are literally the best travel agent in the world and you need to design the best possible trip for your customers"

    # Send the prompt to OpenAI API
    client = OpenAI(api_key=openai_api_key)
    response = client.chat.completions.create(
        model=get_chatbot_model(),
        messages=[{"role": "system", "content": prompt}]
    )
    return response.choices[0].message.content

# Function to generate content using GPT-4
def generate_content_with_gpt4(assistant_prompt, user_prompt):
    client = OpenAI(api_key=openai_api_key)
    messages = [
        {"role": "assistant", "content": assistant_prompt},
        {"role": "user", "content": user_prompt}
    ]
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=messages,
            temperature=0.2,
            max_tokens=256,
            frequency_penalty=0.0
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return str(e)
        
# Function to create Word document for travel guide
def create_word_document_for_travel_guide(toc, chapters):
    doc = Document()
    for i, title in enumerate(toc):
        doc.add_heading(title, level=1)
        doc.add_paragraph(chapters[i])
    return doc
    
# Page 5: Personalized Travel Guide Generator
def page5():
    st.title("📘 Personalized Travel Guide Generator")

    if 'selected_countries' in st.session_state and 'participants' in st.session_state:
        assistant_prompt = "You are a professional travel guide writer."
        user_prompt = f"Please create a table of contents for a travel guide based on the following details: " \
                      f"Countries: {', '.join(st.session_state['selected_countries'])}. " \
                      f"Travelers: {len(st.session_state['participants'])} - "
        for p in st.session_state['participants']:
            user_prompt += f"{p['name']} ({p['age']} years old, {p['gender']}, {p['preference']}). "

        if st.button('Generate Guide'):
            toc = generate_content_with_gpt4(assistant_prompt, user_prompt).split('\n')
            chapters = [generate_content_with_gpt4(assistant_prompt, f"Write a chapter about '{title}' for a travel guide.") for title in toc]
            document = create_word_document_for_travel_guide(toc, chapters)

            # Saving document to a byte stream for download
            doc_stream = io.BytesIO()
            document.save(doc_stream)
            doc_stream.seek(0)
            st.download_button("Download Travel Guide", data=doc_stream, file_name="personalized_travel_guide.docx")
    else:
        st.write("Please complete the previous steps to generate your travel guide.")


# Extend the main function to include page 5
def main():
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ('Page 1', 'Page 2', 'Page 3', 'Page 4', 'Page 5'))

    if page == 'Page 1':
        page1()
    elif page == 'Page 2':
        page2()
    elif page == 'Page 3':
        page3()
    elif page == 'Page 4':
        page4()
    elif page == 'Page 5':
        page5()

if __name__ == "__main__":
    main()

